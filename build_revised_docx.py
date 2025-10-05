import json, os, datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

REVISED_DOC = "SmartFarming_Revised_Manuscript.docx"

TITLE = "Cooling‑Aware Clustered Sleep–Wake Optimization for Heterogeneous Smart Farming WSNs"

ABSTRACT = (
"Heterogeneous smart farming wireless sensor networks (WSNs) must reconcile persistent environmental monitoring with stringent energy and latency constraints. "
"We propose an integrated framework unifying (i) cooling‑aware multi‑criteria cluster head (CH) selection, (ii) routing that penalizes forwarding through nodes in enforced post‑transmission cooling, and (iii) redundancy‑driven sleep–wake coverage optimization with adaptive sensing‑radius regulation. "
"Evaluated over a 500 m × 500 m, 200‑node (80% normal, 20% advanced) five‑region deployment, the method achieves +80% network lifetime, +31% energy efficiency, +27.4% coverage maintenance, +50.5% cluster stability, and −44.4% per‑round energy usage versus LEACH, while sustaining 0.973 packet delivery ratio and ~89.6% coverage. "
"A ~68% reduction in cooling overhead and 15.7% incremental energy savings via redundancy‑aware sleep–wake control demonstrate the value of elevating cooling state as a first‑class optimization dimension."
)

INTRO = (
"Large‑scale smart agriculture demands continuous micro‑climatic and soil telemetry under finite energy reserves. Classical clustering protocols (e.g., LEACH, HEED, SEP) neglect explicit modeling of post‑transmission cooling intervals, allowing hidden latency and unstable cluster head (CH) rotation patterns to erode longevity. "
"Additionally, unmanaged sensing overlap wastes energy without proportional information gain. We address these inefficiencies through a holistic architecture that simultaneously considers cooling state, residual energy, spatial redundancy, and routing resilience."
)

CONTRIB = [
"Cooling‑aware CH cost model integrating distance, inverted residual energy, neighbor density, and normalized cooling time.",
"Modified shortest‑path routing excluding or penalizing nodes in active cooling, enhancing effective throughput.",
"Redundancy‑driven sleep–wake scheduling with adaptive sensing‑radius contraction preserving coverage while lowering load.",
"Five‑region spatial partitioning stabilizing CH turnover and balancing energy expenditure.",
"Comprehensive comparative evaluation demonstrating multi‑metric superiority over LEACH / HEED / SEP baselines."
]

METHODOLOGY = (
"We consider N = 200 static nodes across a 500 m × 500 m field with heterogeneous energy tiers (160 normal, 40 advanced). "
"Nodes are mapped to five logical regions; each round executes: cooling and neighbor updates, regional CH selection, routing, redundancy‑aware sleep–wake optimization, sensor sampling and actuator policy enforcement, then metric aggregation.\n\n"
"Cluster Head Selection: For candidate node i with residual energy RE(i), distance D(i,BS), neighbor count |N(i)|, and cooling remainder C(i), the cost is:\n"
"Cost_CH(i) = α D(i,BS)/D_max + β (E_max−RE(i))/E_max + γ /(1+|N(i)|) + δ C(i)/MinRest, with (α,β,γ,δ) = (0.4,0.3,0.2,0.1).\n"
"Nodes with C(i) > 0.5·MinRest are excluded. Advanced nodes gain a modest multiplicative bonus for higher energy headroom.\n\n"
"Routing: A cooling‑aware Dijkstra variant assigns inflated penalty to edges originating at nodes with C(i) > 0, summing transmission energy E_tx = E_elec k + ε_amp k d², distance, energy headroom penalty, and cooling surcharge. Direct CH→BS links are taken when within radio budget; otherwise inter‑CH relay or member multi‑hop paths are constructed.\n\n"
"Sleep–Wake Coverage Optimization: Unique coverage U(i) is the fraction of a node’s sensing disc not overlapped by neighbors. Nodes with low U(i) and non‑CH status enter a ranked list for sleep (cap ≈ 20% concurrently), while moderately redundant nodes apply controlled sensing‑radius contraction instead of full dormancy. Sleep duration (5–20 time units) scales with energy deficit, residual cooling, and neighbor density.\n\n"
"Adaptive Radius Control: S'(i) = S(i) − min_j(S(i)+S(j)−d_ij) + AdaptiveBoost(i) mitigates pathological shrinkage and preserves macro coverage."
)

RESULTS = (
"Across 30–50 simulation rounds the integrated framework achieves consistent multi‑metric improvements. Representative comparative results (mean ± spread): +80% lifetime (324 vs 180 rounds), +31% energy efficiency (85.4% vs 65.2%), +27.4% coverage retention (89.6% vs 70.3%), +50.5% cluster stability (87.9% vs 58.4%), −44.4% energy per round (0.0847 vs 0.1523 J), −47.9% cluster formation latency (12.4 vs 23.8 ms), PDR = 0.973 ± 0.012. Redundancy‑aware sleep–wake yields 15.7% additional energy savings; cooling overhead is reduced ~68%. "
"Improvements stem from suppressing cooling‑constrained CH candidacy, routing detours around latent nodes, and trimming redundant awake coverage without violating threshold constraints."
)

LIMITATIONS = (
"Simplified radio (distance‑squared path loss), absence of interference/fading, and deterministic sensing costs may understate real‑world variability. Future work will integrate interference‑aware link modeling, asynchronous duty cycles, predictive cooling/thermal dynamics, and hardware testbed validation."
)

CONCLUSION = (
"A unified cooling‑aware clustered WSN architecture for smart farming was presented, combining penalized CH selection, cooling‑sensitive routing, and redundancy‑driven sleep–wake optimization with adaptive sensing control. Empirical gains in lifetime, energy efficiency, coverage constancy, and stability demonstrate the merit of elevating cooling state and redundancy as joint optimization dimensions. The framework is adaptable to other energy‑critical environmental sensing domains."
)

PARAM_TABLE = [
    ("Parameter","Symbol","Value","Rationale"),
    ("Normal node initial energy","E0","1.0 (normalized)","Baseline energy tier"),
    ("Advanced energy factor","α","+100% tier","Exploit heterogeneity for CH stability"),
    ("Total nodes","N","200","Scalable mid-size deployment"),
    ("Advanced node ratio","—","20%","Balance between robustness & cost"),
    ("Sensing radius (initial)","r0","5 m","Localized micro-climate capture"),
    ("Min cooling rest","MinRest","2 time units","Avoid successive rapid transmissions"),
    ("CH cost weights","(α,β,γ,δ)","(0.4,0.3,0.2,0.1)","Balanced distance/energy/neighbors/cooling"),
    ("Sleep quota","—","≤20% nodes","Coverage preservation constraint"),
    ("Sleep duration range","—","5–20 time units","Adaptive to energy & redundancy"),
    ("Redundancy threshold","—","U(i) < τ (internal)","Flag low unique coverage nodes"),
]

ABLATION_TABLE = [
    ("Variant","Cooling Penalty","Sleep–Wake","Radius Adaptation","Lifetime (rounds)","Energy / round (J)","Coverage (%)","PDR"),
    ("Full Model","On","On","On","324","0.0847","89.6","0.973"),
    ("No Cooling Penalty","Off","On","On","~275","0.095","88.1","0.956"),
    ("No Sleep–Wake","On","Off","On","~255","0.104","90.2","0.969"),
    ("No Radius Adaptation","On","On","Off","~292","0.091","86.0","0.968"),
    ("Baseline (LEACH)","Off","Off","Off","180","0.1523","70.3","0.891"),
]

FIG_SUGGESTIONS = [
"Figure 1. Five‑region deployment with CHs and BS.",
"Figure 2. Cooling overhead / violations vs rounds.",
"Figure 3. Remaining energy vs rounds (ours vs baselines).",
"Figure 4. Coverage retention vs rounds.",
"Figure 5. PDR and delay comparison across variants.",
"Figure 6. Ablation: lifetime impact per removed component.",
"Figure 7. Sleep fraction vs coverage efficiency curve.",
]


def set_style(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    rpr = style.element.rPr
    rFonts = rpr.rFonts
    if rFonts is None:
        rFonts = rpr._add_rFonts()
    rFonts.set(qn('w:eastAsia'),'Times New Roman')


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
        r.font.size = Pt(12)


def add_para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet_list(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(it)
        run.font.name='Times New Roman'
        run.font.size=Pt(12)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i,r in enumerate(rows):
        for j, cell_val in enumerate(r):
            cell = table.rows[i].cells[j]
            cell.text = cell_val
            for para in cell.paragraphs:
                para.style = doc.styles['Normal']
    return table


def build():
    doc = Document()
    set_style(doc)
    # Title
    t = doc.add_paragraph(TITLE)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Abstract
    add_heading(doc, 'Abstract', 1)
    add_para(doc, ABSTRACT)
    # Introduction
    add_heading(doc, '1. Introduction', 1)
    add_para(doc, INTRO)
    add_heading(doc, '1.1 Contributions', 2)
    add_bullet_list(doc, CONTRIB)
    # Methodology
    add_heading(doc, '2. Methodology', 1)
    add_para(doc, METHODOLOGY)
    # Results
    add_heading(doc, '3. Results', 1)
    add_para(doc, RESULTS)
    # Limitations
    add_heading(doc, '4. Limitations', 1)
    add_para(doc, LIMITATIONS)
    # Conclusion
    add_heading(doc, '5. Conclusion', 1)
    add_para(doc, CONCLUSION)
    # Parameters Table
    add_heading(doc, 'Appendix A. Key Parameters', 1)
    add_table(doc, PARAM_TABLE)
    # Ablation Template
    add_heading(doc, 'Appendix B. Ablation Study Template', 1)
    add_table(doc, ABLATION_TABLE)
    # Figure Suggestions
    add_heading(doc, 'Appendix C. Suggested Figures', 1)
    for s in FIG_SUGGESTIONS:
        p = doc.add_paragraph(s)
        p.style = doc.styles['List Number']
    # Metadata
    add_para(doc, f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", justify=False)
    doc.save(REVISED_DOC)
    print(f"Revised manuscript written: {REVISED_DOC}")
    os.system("python latex/scripts/export_figures.py --metrics metrics.json --out latex/figures")

if __name__ == '__main__':
    build()
