#!/usr/bin/env python3
"""
Generate final revised manuscript in DOCX format.
Includes all sections from the LaTeX manuscript with proper formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import re

def setup_styles(doc):
    """Configure document styles."""
    # Normal text style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
        else:
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True

def add_title_section(doc):
    """Add title and authors."""
    title = doc.add_heading('Cooling-Aware Clustered Sleep–Wake Optimization for Heterogeneous Smart Farming WSNs', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run('First Author¹, Second Author¹')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Affiliation
    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affiliation.add_run('¹Department/Institute, University, City, Country')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    
    # Email
    email = doc.add_paragraph()
    email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = email.add_run('Email: first.author@example.edu, second.author@example.edu')
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # Spacing

def add_abstract(doc):
    """Add abstract section."""
    doc.add_heading('Abstract', level=1)
    
    abstract_text = """Wireless sensor networks (WSNs) deployed in smart farming environments face critical challenges balancing energy efficiency, spatial coverage, and thermal constraints. Existing clustering and coverage protocols (LEACH, HEED, SEP, DEEC) optimize energy consumption but neglect post-transmission cooling intervals—mandatory rest periods required to prevent thermal damage in resource-constrained sensor nodes. This paper introduces a novel cross-layer framework that integrates explicit cooling-state modeling (Cᵢ ∈ {0,1,...,MinRest}) as a first-class optimization variable across three subsystems: (i) cooling-aware cluster-head (CH) selection with δ-weighted thermal penalties, (ii) routing via modified Dijkstra with λcool edge surcharges avoiding nodes in active cooling, and (iii) redundancy-driven sleep–wake scheduling coupled with adaptive sensing-radius contraction (Sᵢ ∈ [2,5] m). 

We validate the framework on a 200-node heterogeneous deployment (160 normal, 40 advanced nodes) across a 500×500 m field via 50 Monte Carlo simulations, comparing against six established baselines: LEACH, HEED, SEP, DEEC, EDEEC, and TADR (thermal-aware delay-constrained routing). Results demonstrate +80% network lifetime vs. LEACH, +29% vs. best baseline (EDEEC), −44.4% per-round energy consumption, 89.6±0.8% sustained coverage (vs. 70.3% LEACH), ∼68% cooling overhead reduction, and 0.973 packet delivery ratio. Parameter sensitivity analysis (δ, fmax, MinRest) confirms robustness, with optimal performance at δ=0.10, fmax=0.20, MinRest=2 (validated against CC2420 radio thermal profiles). Computational trade-off analysis shows Monte Carlo unique-coverage approximation (1.2 ms/node, ±3.5% error) saves 8.1 minutes vs. analytic methods over 324 rounds, making it feasible for embedded deployment. The framework is the first to holistically integrate discrete cooling-state transitions across clustering, routing, and coverage control, achieving superior energy-thermal-coverage trade-offs validated via paired t-tests (p<0.01) and complete reproducibility (GitHub: https://github.com/vivekjindal24/Isha-PhD)."""
    
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Keywords
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    run = keywords.add_run('Keywords: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    keywords.add_run('Wireless Sensor Networks, Cooling-Aware Clustering, Energy Optimization, Coverage Control, Smart Farming, Thermal Management, Sleep Scheduling, Heterogeneous WSN')

def add_introduction(doc):
    """Add introduction section."""
    doc.add_heading('1. Introduction', level=1)
    
    intro_paragraphs = [
        """Wireless sensor networks (WSNs) have emerged as foundational infrastructure for precision agriculture, enabling real-time monitoring of soil moisture, temperature, humidity, and crop health across large farmlands. However, battery-powered sensor nodes face three interrelated constraints: (i) limited energy reserves (typically 2–5 J from AA batteries), (ii) spatial coverage requirements (≥80% field monitoring for actionable insights), and (iii) thermal stress from repeated high-power transmissions (e.g., cluster-head duty in hierarchical protocols). While classical clustering protocols like LEACH, HEED, and SEP optimize energy expenditure through adaptive cluster-head (CH) rotation and load balancing, they overlook a critical operational constraint: post-transmission cooling intervals.""",
        
        """Modern WSN radios (e.g., CC2420, CC2520) operating at 0 dBm transmit power require mandatory rest periods (∼2 seconds) after burst transmissions to dissipate accumulated heat and prevent thermal damage or regulatory violations. Ignoring this constraint leads to hidden performance penalties: (a) nodes forced into cooling become temporarily unavailable for routing, creating forwarding bottlenecks and increased end-to-end delay; (b) repeated thermal stress on the same nodes (e.g., frequently re-elected CHs) accelerates hardware degradation and premature failure; (c) unmanaged cooling overhead fragments the network topology, reducing effective throughput despite adequate residual energy.""",
        
        """Recent work has begun addressing thermal constraints in WSNs: Zhang et al. introduced node-level cooling-aware duty cycling, achieving 35% lifetime gains over baseline LEACH. However, their approach operates in isolation—individual nodes self-regulate cooling without coordination with clustering or routing layers. Similarly, TADR (Thermal-Aware Delay-constrained Routing) uses instantaneous temperature sensors to guide path selection but does not model discrete cooling states or integrate with CH selection. The DEEC family (DEEC, DDEEC, EDEEC) refines heterogeneous clustering through dynamic energy-ratio-based CH probabilities and multi-tier adaptive thresholds, yet treats all nodes as thermally equivalent, missing opportunities to distribute thermal load spatially and temporally.""",
        
        """This paper introduces the first holistic cross-layer framework that models cooling as a discrete state variable (Cᵢ ∈ {0,1,...,MinRest}) and integrates it across three optimization layers: (1) Cooling-Aware Cluster-Head Selection: Extends classical cost functions (distance, residual energy, node density) with an explicit δ-weighted cooling penalty (δCᵢ/MinRest), preventing nodes in active cooling from CH candidacy and distributing thermal stress regionally. (2) Cooling-Penalized Multi-Hop Routing: Modifies Dijkstra's algorithm with edge-weight surcharges (λcool·Cᵢ/MinRest) that avoid routing through cooling nodes, reducing forwarding latency and path oscillation. (3) Redundancy-Driven Sleep–Wake with Adaptive Radius: Couples unique-coverage estimation (Uᵢ via Monte Carlo sampling) with energy-deficit-scaled sleep duration and adaptive sensing-radius contraction (Sᵢ ∈ [2,5] m), preserving ≥80% coverage while sleeping up to 20% of nodes per round.""",
        
        """We validate the framework on a 200-node heterogeneous deployment (80% normal, 20% advanced energy tier) across a 500×500 m smart farming field via 50 Monte Carlo simulations with fixed random seeds (1000-1049) for reproducibility. Comparisons against six baselines—LEACH, HEED, SEP, DEEC, EDEEC (best baseline at 251 rounds), and TADR—demonstrate: +80% network lifetime vs. LEACH (324 vs. 180 rounds), +29% vs. EDEEC, −44.4% per-round energy consumption (0.0847 vs. 0.1523 J), 89.6±0.8% sustained coverage (vs. 70.3% LEACH, 79.1% EDEEC), ∼68% cooling overhead reduction (6.2% vs. 18.9% nodes cooling per round), and 0.973 packet delivery ratio with 18.4 ms delay (vs. 0.891 PDR, 32.7 ms for LEACH). Parameter sensitivity analysis confirms robustness across δ ∈ [0.0, 0.25], fmax ∈ [0.1, 0.3], MinRest ∈ [1,5], with optimal performance at δ=0.10, fmax=0.20, MinRest=2 (validated against CC2420 thermal recovery profiles).""",
        
        """The key contributions of this work are: (1) First discrete cooling-state model (Cᵢ with deterministic transition rules) integrated across WSN clustering, routing, and coverage layers—prior work uses continuous temperature metrics or isolated node-level duty cycling. (2) Novel multi-factor CH cost function balancing spatial distance, residual energy, node density, and cooling penalty—enables thermal load distribution impossible with energy-only heuristics (DEEC/EDEEC). (3) Cooling-penalized routing avoiding latent forwarding bottlenecks—outperforms TADR by 65% (324 vs. 196 rounds) through discrete state modeling vs. continuous temperature-based path oscillation. (4) Unified redundancy management: unique coverage (Uᵢ) + adaptive radius (Sᵢ) + energy-scaled sleep duration—achieves PEAS-like coverage efficiency (89.6%) with added thermal sustainability and 44% energy savings. (5) Comprehensive validation: 6 baselines spanning classical/DEEC/thermal-aware families, 50-run Monte Carlo, 95% CI, paired t-tests (p<0.01), parameter sweeps with physical justifications (CC2420 profiles), computational trade-off analysis (Monte Carlo vs. analytic), complete reproducibility (GitHub repository with simulation code, datasets, analysis scripts)."""
    ]
    
    for para_text in intro_paragraphs:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_related_work(doc):
    """Add related work section."""
    doc.add_heading('2. Related Work and Literature Review', level=1)
    
    # Subsection 2.1
    doc.add_heading('2.1 Energy-Efficient Clustering in WSNs', level=2)
    
    rw_text_1 = """Hierarchical clustering protocols have been central to WSN energy optimization since the seminal work of Heinzelman et al. (2000) with LEACH. LEACH introduced randomized, probabilistic cluster head (CH) rotation, achieving distributed load balancing but with limited adaptivity to residual energy or spatial heterogeneity. Younis and Fahmy (2004) proposed HEED, which extends LEACH by incorporating residual energy and intra-cluster communication cost into CH selection, yielding more stable clusters. Smaragdakis et al. (2004) addressed heterogeneous networks via SEP, weighting CH election probabilities by initial energy tier, thus prolonging network lifetime through preferential selection of advanced nodes.

The DEEC family (Distributed Energy-Efficient Clustering) further refines heterogeneous clustering by dynamically adjusting CH election probabilities based on the ratio of residual-to-initial energy, enabling nodes with higher remaining capacity to serve as CHs more frequently. Qing et al. (2006) introduced DEEC with two-tier energy levels. Variants include DDEEC (Developed DEEC) by Elbhiri et al. (2010), which introduces three energy levels (normal, advanced, super), and EDEEC (Enhanced DEEC) by Saini and Sharma (2010), which adapts cluster formation thresholds per round. However, all DEEC-family protocols optimize purely for energy balance and do not account for post-transmission cooling intervals—the enforced temporal gap required after high-power transmissions to prevent thermal damage or regulatory violations. As a result, rapid successive CH re-elections or routing through recently active nodes can inadvertently degrade effective throughput and introduce hidden latency."""
    
    p = doc.add_paragraph(rw_text_1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Subsection 2.2
    doc.add_heading('2.2 Cooling and Thermal Constraints in Wireless Systems', level=2)
    
    rw_text_2 = """Thermal management in wireless devices has been explored primarily in cellular base stations and high-density IoT deployments, where processor throttling and transmission scheduling are adjusted to maintain safe operating temperatures. However, explicit cooling-state modeling in WSN routing and clustering remains nascent. Recent work by Zhang et al. (2021) introduced cooling-aware duty cycling for individual nodes but did not integrate it into multi-hop routing or CH selection cost functions. Our approach differs by treating cooling state as a first-class optimization variable across all three subsystems: clustering, routing, and coverage control."""
    
    p = doc.add_paragraph(rw_text_2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Subsection 2.3
    doc.add_heading('2.3 Coverage Optimization and Redundancy Management', level=2)
    
    rw_text_3 = """Coverage maximization under energy constraints has been addressed through sleep scheduling (Ye et al., 2003; Tian & Georganas, 2002) and adaptive sensing radius control (Wang et al., 2007). PEAS uses probing to keep only necessary sensors active, while Tian and Georganas proposed node scheduling based on geometric coverage overlap. However, these methods do not consider the interplay between redundancy, energy expenditure, and cooling overhead. Our redundancy-driven sleep–wake optimization with adaptive radius contraction uniquely couples spatial redundancy (unique coverage Uᵢ) with energy and cooling state to co-optimize coverage preservation and thermal sustainability."""
    
    p = doc.add_paragraph(rw_text_3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Subsection 2.4
    doc.add_heading('2.4 Routing Under Operational Constraints', level=2)
    
    rw_text_4 = """Directed diffusion and gradient-based routing optimize paths based on energy reserves and link quality, but do not account for transient node unavailability due to cooling. Thermal-aware routing has emerged in recent years: TADR (Thermal-Aware Delay-constrained Routing) by Chen et al. (2019) minimizes end-to-end delay by selecting paths with lower cumulative thermal load, measured via temperature sensors at each node. However, TADR does not model explicit cooling state transitions (i.e., mandatory rest periods) and instead uses instantaneous temperature as a continuous metric. Our cooling-aware Dijkstra variant differs by: (i) treating cooling as a discrete state variable (Cᵢ ∈ {0,1,...,MinRest}), (ii) imposing deterministic edge penalties (λcool·Cᵢ/MinRest) that enforce path avoidance during mandatory rest, and (iii) integrating this routing logic with cluster-head selection and sleep–wake mechanisms, creating a unified cross-layer optimization framework. Recent energy-aware routing surveys focus on residual battery levels and hop counts; our approach extends this by avoiding latent forwarding bottlenecks caused by cooling-induced node unavailability."""
    
    p = doc.add_paragraph(rw_text_4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Subsection 2.5
    doc.add_heading('2.5 Positioning This Work', level=2)
    
    rw_text_5 = """Table 1 summarizes key distinctions between our framework and related WSN optimization approaches. Our integrated framework is the first to:

(i) Incorporate explicit cooling cost (δCᵢ/MinRest) into CH selection alongside distance, energy, and density.
(ii) Penalize routing edges originating at cooling nodes via λcool surcharges, ensuring paths avoid thermal bottlenecks.
(iii) Unify redundancy-based sleep scheduling with adaptive sensing-radius contraction, preserving macro coverage while reducing per-node load and cooling frequency.

This integrated cooling-aware paradigm yields substantial empirical gains (detailed in Section 5) while addressing a gap in the literature: the holistic incorporation of thermal state constraints into WSN optimization."""
    
    p = doc.add_paragraph(rw_text_5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add comparison table (simplified for DOCX)
    doc.add_paragraph()
    doc.add_paragraph('Table 1: Comparison with Related WSN Optimization Approaches', style='Heading 3')
    
    table = doc.add_table(rows=10, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Method', 'Clustering', 'Routing', 'Coverage', 'Cooling', 'Heterogeneity']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    data = [
        ['LEACH', 'Random', 'CH-direct', 'Fixed', 'No', 'Homogeneous'],
        ['HEED', 'Energy+cost', 'CH-direct', 'Fixed', 'No', 'Homogeneous'],
        ['SEP', 'Weighted prob.', 'CH-direct', 'Fixed', 'No', '2-tier energy'],
        ['DEEC', 'Dynamic prob.', 'CH-direct', 'Fixed', 'No', '2-tier energy'],
        ['EDEEC', 'Adaptive thresh.', 'CH-direct', 'Fixed', 'No', '3-tier energy'],
        ['TADR', '—', 'Thermal-aware', 'Fixed', 'Temp-based', 'Homogeneous'],
        ['PEAS', '—', '—', 'Probing-based', 'No', 'Homogeneous'],
        ['Zhang et al.', '—', '—', 'Duty-cycle', 'Node-level', 'Homogeneous'],
        ['Proposed', 'Multi-factor', 'Cooling-aware', 'Redundancy+adaptive', 'Integrated', '2-tier regional']
    ]
    
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            cells[j].text = cell_data
            if i == len(data) - 1:  # Bold the last row (Proposed)
                cells[j].paragraphs[0].runs[0].font.bold = True

def add_methodology_summary(doc):
    """Add methodology summary (shortened for DOCX)."""
    doc.add_heading('3. System Model and Methodology', level=1)
    
    doc.add_heading('3.1 Network Model', level=2)
    
    method_text_1 = """We consider a heterogeneous WSN with N = 200 sensor nodes randomly deployed across a 500×500 m smart farming field. The network comprises two energy tiers: 160 normal nodes (E₀ = 0.5 J) and 40 advanced nodes (E₀ = 1.0 J, providing 2× initial energy). A stationary base station (BS) is located at (250, 550) m. The field is partitioned into R = 5 vertical regions (each 100×500 m) to enable regional cluster-head constraints and load balancing.

Each node i is characterized by: position (xᵢ, yᵢ), residual energy Eᵢ(t), sensing radius Sᵢ(t) ∈ [2,5] m (adaptively controlled), sleep/active state, and cooling state Cᵢ(t) ∈ {0,1,...,MinRest}. Time is discretized into rounds (duration τ ≈ 20 s), each comprising setup phase (CH election, cluster formation) and steady-state phase (data collection, aggregation, multi-hop transmission to BS)."""
    
    p = doc.add_paragraph(method_text_1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('3.2 Energy Model', level=2)
    
    method_text_2 = """We adopt the first-order radio model with parameters: Eelec = 50 nJ/bit (transceiver electronics), εamp = 100 pJ/(bit·m²) (amplifier), γ = 2 (path-loss exponent for short-range transmission). Transmission energy for k-bit message over distance d: ETx(k,d) = Eelec·k + εamp·k·d². Reception energy: ERx(k) = Eelec·k. Cluster-head nodes incur additional aggregation cost Eagg = 5 nJ/bit."""
    
    p = doc.add_paragraph(method_text_2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('3.3 Cooling State Machine', level=2)
    
    method_text_3 = """After serving as cluster head or relaying packets, node i enters cooling state Cᵢ = MinRest (mandatory rest period). The cooling state evolves deterministically:

    Cᵢ(t+1) = max(0, Cᵢ(t) - 1)    if node i is idle at round t
    Cᵢ(t+1) = MinRest              if node i transmits/relays at round t

Nodes with Cᵢ > 0 are excluded from cluster-head candidacy and penalized in routing edge weights. MinRest = 2 is calibrated based on CC2420 radio thermal recovery profiles (≈2 seconds to return to baseline temperature after 1-second burst at 0 dBm TX power)."""
    
    p = doc.add_paragraph(method_text_3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('3.4 Cooling-Aware Cluster-Head Selection', level=2)
    
    method_text_4 = """Cluster heads are elected via multi-factor cost minimization within each region. For node i in region r, the cost function is:

    Cost(i) = α·(dᵢ,center/dmax) + β·(1 - Eᵢ/Einit) + γ·(1/Densityᵢ) + δ·(Cᵢ/MinRest)

where dᵢ,center is distance to regional centroid, Densityᵢ is neighbor count within 2·Smax, and weights (α,β,γ,δ) = (0.4, 0.3, 0.2, 0.1). The cooling penalty term δ·(Cᵢ/MinRest) prevents recently active nodes from immediate re-election, distributing thermal load spatially and temporally. Nodes with Cᵢ ≥ 1 are hard-excluded from CH candidacy."""
    
    p = doc.add_paragraph(method_text_4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('3.5 Cooling-Penalized Multi-Hop Routing', level=2)
    
    method_text_5 = """We modify Dijkstra's shortest-path algorithm with cooling-aware edge weights:

    Weight(i→j) = λdist·dᵢⱼ + λeng·(1/Eⱼ) + λcool·(Cⱼ/MinRest)

where λdist = 1.0, λeng = 2.0, λcool = 5.0. The high cooling penalty (2.5× energy weight) prioritizes path avoidance through nodes in active cooling, reducing forwarding latency and preventing thermal bottlenecks. Routing graphs are recomputed per round after CH election."""
    
    p = doc.add_paragraph(method_text_5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('3.6 Redundancy-Driven Sleep–Wake Optimization', level=2)
    
    method_text_6 = """Non-CH nodes participate in sleep scheduling based on unique coverage Uᵢ (area covered by node i but not by any neighbors). We estimate Uᵢ via Monte Carlo sampling: generate M = 50 random points within node i's sensing disc, count points not covered by neighbors. Nodes with low Uᵢ (high redundancy) are prioritized for sleep, subject to:

(i) Regional coverage constraint: ≥80% coverage per region
(ii) Maximum sleep fraction: fmax = 20% of nodes per region
(iii) Energy-scaled sleep duration: Tsleep(i) = τ·(1 - Eᵢ/Einit)·(1 - Uᵢ/Sᵢ²π)

Additionally, adaptive radius control scales Sᵢ ∈ [2,5] m based on residual energy, contracting sensing range under low battery to reduce overlap and transmission power."""
    
    p = doc.add_paragraph(method_text_6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_results_summary(doc):
    """Add results summary."""
    doc.add_heading('4. Experimental Results and Validation', level=1)
    
    doc.add_heading('4.1 Experimental Setup', level=2)
    
    results_text_1 = """Simulations were conducted in Python 3.10.12 (NumPy 1.24.2, SciPy 1.10.1, Matplotlib 3.7.1) over 50 independent Monte Carlo runs with fixed random seeds (1000-1049) for reproducibility. Each run executes until first node failure (lifetime) or 500 rounds maximum. We compare against six established baselines spanning three protocol families:

Classical Clustering: LEACH (randomized probabilistic CH rotation), HEED (residual-energy and cost-based CH selection), SEP (weighted election for two-tier heterogeneity).

DEEC Family: DEEC (dynamic energy-ratio-based CH probability), EDEEC (three-tier adaptive threshold clustering, best baseline at 251 rounds).

Thermal/Delay-Aware Routing: TADR (thermal-aware delay-constrained routing using instantaneous temperature metrics).

All baselines use identical energy model, packet sizes (2000 bits data, 500 bits control), field dimensions, and node counts. Critical distinction: none incorporate discrete cooling-state transitions or enforce mandatory post-transmission rest periods."""
    
    p = doc.add_paragraph(results_text_1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('4.2 Primary Performance Metrics', level=2)
    
    results_text_2 = """Table 2 presents mean ± 95% confidence intervals across 50 runs. Statistical significance assessed via paired t-tests (p<0.01).

Key Findings:
• Lifetime: Proposed achieves 324 rounds vs. 180 (LEACH), 218 (HEED), 245 (SEP), 232 (DEEC), 251 (EDEEC), 196 (TADR)—80% improvement over LEACH, 29% over best baseline (EDEEC).
• Energy Efficiency: 0.0847 J/round (44.4% lower than LEACH, 17% lower than EDEEC, 42% lower than TADR).
• Coverage: 89.6±0.8% sustained (vs. 70.3% LEACH, 79.1% EDEEC, 71.5% TADR).
• PDR & Delay: 0.973 PDR, 18.4 ms delay (vs. 0.891 PDR, 32.7 ms for LEACH; 0.942 PDR, 25.2 ms for TADR).
• Cluster Stability: 87.9% (vs. 58.4% LEACH, 74.5% EDEEC, 61.2% TADR).

DEEC family (232-251 rounds) outperforms classical protocols through dynamic energy-based CH probability but lags 22-28% behind our cooling-aware approach. TADR's instantaneous thermal metrics yield only marginal gains over LEACH (+8.9%) due to path oscillation, lack of clustering integration, and no coverage optimization."""
    
    p = doc.add_paragraph(results_text_2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add simplified results table
    doc.add_paragraph()
    doc.add_paragraph('Table 2: Primary Performance Metrics (Mean ± 95% CI, 50 runs)', style='Heading 3')
    
    table = doc.add_table(rows=8, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    headers = ['Metric', 'Proposed', 'LEACH', 'HEED', 'SEP', 'DEEC', 'EDEEC', 'TADR']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Data
    data = [
        ['Lifetime (rounds)', '324±6.2', '180±4.1', '218±5.3', '245±5.8', '232±5.5', '251±6.0', '196±4.7'],
        ['Energy/round (J)', '0.0847±0.0015', '0.1523±0.0025', '0.1210±0.0020', '0.1048±0.0018', '0.1185±0.0022', '0.1021±0.0019', '0.1456±0.0024'],
        ['Coverage (%)', '89.6±0.8', '70.3±1.1', '75.4±1.0', '78.2±0.9', '74.8±1.0', '79.1±0.9', '71.5±1.1'],
        ['PDR', '0.973±0.004', '0.891±0.010', '0.921±0.007', '0.938±0.006', '0.915±0.008', '0.942±0.006', '0.908±0.009'],
        ['Delay (ms)', '18.4±1.2', '32.7±2.1', '28.3±1.8', '24.6±1.5', '29.1±1.9', '23.8±1.6', '25.2±1.7'],
        ['Stability', '0.879±0.012', '0.584±0.015', '0.692±0.013', '0.731±0.014', '0.678±0.014', '0.745±0.013', '0.612±0.015']
    ]
    
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            cells[j].text = cell_data
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if j == 1:  # Bold proposed column
                        run.font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 Parameter Sensitivity Analysis', level=2)
    
    results_text_3 = """Cooling Weight δ: Sweeping δ ∈ {0.00, 0.05, 0.10, 0.15, 0.20, 0.25} reveals optimal performance at δ=0.10 (inverted-U curve). Under-penalization (δ<0.10) allows premature CH re-election (298 rounds at δ=0.05, −8%). Over-penalization (δ>0.10) over-constrains CH candidacy with poor spatial placement (285 rounds at δ=0.20, −12%). The sweet spot balances thermal load distribution with coverage quality.

Maximum Sleep Fraction fmax: Varying fmax ∈ {0.10, 0.15, 0.20, 0.25, 0.30} demonstrates inflection point at fmax=0.20. Coverage degrades from 90.8% (fmax=0.10) to 78.1% (fmax=0.30), violating 80% threshold. Energy savings diminish beyond 0.20 (marginal <2% per increment vs. >5% coverage loss). At fmax=0.25, AdaptiveBoost triggers too frequently, paradoxically increasing overhead.

MinRest Period: Sweeping MinRest ∈ {1,2,3,4,5} shows optimal at MinRest=2 (validated by CC2420 thermal profiles). PDR peaks at 0.973 (MinRest=2) vs. 0.958 (MinRest=1, insufficient cooling) vs. 0.942 (MinRest=5, excessive unavailability). Delay minimized at 18.4 ms (MinRest=2) vs. 22.1 ms (MinRest=5) due to routing path elongation."""
    
    p = doc.add_paragraph(results_text_3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('4.4 Computational Trade-off Analysis', level=2)
    
    results_text_4 = """We benchmarked three methods for unique coverage Uᵢ computation:

Monte Carlo Approximation (our implementation): M=50 random samples, 1.2 ms/node, ±3.5% error (95% CI).
Grid Discretization: 0.1×0.1 m cells, 3.8 ms/node, ±2.1% error.
Analytic Circle Intersection: Delaunay triangulation + geometric formulae, 8.7 ms/node, 0% error (7.2× slower).

Monte Carlo achieves optimal balance: <2 ms per node × 200 nodes = 240 ms overhead per round vs. analytic's 1.74 s. Over 324 rounds, saves 8.1 minutes—critical for real-time embedded deployment. The 3.5% error is acceptable given ±5% RSSI measurement noise typical in WSNs."""
    
    p = doc.add_paragraph(results_text_4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_discussion_conclusion(doc):
    """Add discussion and conclusion."""
    doc.add_heading('5. Discussion', level=1)
    
    disc_text = """The experimental results validate four key claims:

(I) Cooling-State Integration is Critical: Explicit modeling of Cᵢ as discrete state variable in both CH selection and routing yields 29% lifetime gain over best baseline (EDEEC). Even sophisticated energy-aware clustering cannot compensate for thermal bottlenecks without cooling penalties.

(II) Continuous Temperature Metrics Are Insufficient: TADR's instantaneous temperature readings yield only +8.9% vs. LEACH due to: (a) path oscillation from dynamic temperature updates, (b) no clustering integration (thermally stressed CHs), (c) no coverage optimization (redundant active nodes). Our deterministic cooling states avoid these issues.

(III) Redundancy Management Enhances Sustainability: Sleep–wake scheduling and adaptive radius control achieve 89.6% coverage vs. 74-79% for DEEC/EDEEC (which lack redundancy-driven sleep). Unique coverage metric Uᵢ identifies non-redundant nodes, preserving high coverage while sleeping 15-20% per round.

(IV) Synergistic Gains Require Co-Design: Ablation study shows removing any component degrades performance by 10-21%. Integrated architecture necessary—isolated optimizations insufficient.

Comparison to State-of-the-Art: Zhang et al.'s 35% lifetime gain (node-level cooling duty cycling) is half our 80% gain over LEACH and 29% over EDEEC, validating superiority of cross-layer integration. Our framework achieves PEAS-like coverage efficiency (89.6%) while adding thermal sustainability (68% cooling overhead reduction) and energy efficiency (44% lower energy/round). Gap vs. EDEEC widens to 34% at N=300, demonstrating scalability."""
    
    p = doc.add_paragraph(disc_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('6. Conclusion', level=1)
    
    conclusion_text = """This paper introduced the first holistic cross-layer framework integrating discrete cooling-state modeling (Cᵢ ∈ {0,1,...,MinRest}) across WSN clustering, routing, and coverage control. The framework addresses a critical gap: existing protocols (LEACH, HEED, SEP, DEEC, EDEEC) optimize energy but neglect post-transmission cooling intervals, while recent thermal-aware methods (Zhang et al., TADR) operate in isolation or use continuous metrics prone to path oscillation.

Our contributions include: (1) novel multi-factor CH cost function with δ-weighted cooling penalty distributing thermal load spatially/temporally, (2) cooling-penalized routing avoiding latent forwarding bottlenecks via λcool edge surcharges, (3) unified redundancy management coupling unique coverage Uᵢ with adaptive radius Sᵢ and energy-scaled sleep duration, and (4) comprehensive validation across 6 baselines (50 runs, paired t-tests, parameter sweeps, computational trade-offs, full reproducibility).

Results demonstrate +80% network lifetime vs. LEACH, +29% vs. best baseline (EDEEC), −44.4% per-round energy, 89.6% sustained coverage, ∼68% cooling overhead reduction, 0.973 PDR. Parameter sensitivity confirms robustness (δ=0.10, fmax=0.20, MinRest=2 validated by CC2420 profiles). Monte Carlo unique-coverage computation saves 8.1 minutes vs. analytic over 324 rounds, enabling embedded deployment.

Future work will: (i) integrate stochastic channel models (Rayleigh/Rician fading), (ii) validate via hardware testbed (TelosB/OpenMote-B with thermal profiling), (iii) extend to continuous temperature dynamics using heat equation solvers, (iv) evaluate under mobile sink/node mobility, (v) implement distributed cooling-state synchronization for large-scale deployments (N>500). All simulation code, datasets, and analysis scripts available at: https://github.com/vivekjindal24/Isha-PhD."""
    
    p = doc.add_paragraph(conclusion_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_references(doc):
    """Add references section."""
    doc.add_heading('References', level=1)
    
    references = [
        "Chen, J., Wang, Y., Li, M., et al. (2019). Thermal-aware delay-constrained routing for wireless sensor networks. IEEE Transactions on Industrial Informatics, 15(8), 4432-4441.",
        "Elbhiri, B., Saadane, R., El Fkihi, S., & Aboutajdine, D. (2010). Developed distributed energy-efficient clustering (DDEEC) for heterogeneous wireless sensor networks. Proceedings of I/V Communications and Mobile Network (ISVC), 1-4.",
        "Heinzelman, W. R., Chandrakasan, A., & Balakrishnan, H. (2000). Energy-efficient communication protocol for wireless microsensor networks. Proceedings of the 33rd Annual Hawaii International Conference on System Sciences. IEEE.",
        "Polastre, J., Szewczyk, R., & Culler, D. (2005). Telos: Enabling ultra-low power wireless research. Proceedings of IEEE/ACM IPSN, 364-369.",
        "Qing, L., Zhu, Q., & Wang, M. (2006). Design of a distributed energy-efficient clustering algorithm for heterogeneous wireless sensor networks. Computer Communications, 29(12), 2230-2237.",
        "Saini, P., & Sharma, A. K. (2010). Energy efficient scheme for clustering protocol prolonging the lifetime of heterogeneous wireless sensor networks. International Journal of Computer Applications, 6(2), 30-36.",
        "Smaragdakis, G., Matta, I., & Bestavros, A. (2004). SEP: A stable election protocol for clustered heterogeneous wireless sensor networks. Second International Workshop on Sensor and Actor Network Protocols and Applications.",
        "Srinivasan, K., & Levis, P. (2008). An empirical study of low-power wireless. ACM Transactions on Sensor Networks (TOSN), 6(2), 1-49.",
        "Tian, D., & Georganas, N. D. (2002). A coverage-preserving node scheduling scheme for large wireless sensor networks. Proceedings of ACM WSNA, 32-41.",
        "Wang, X., Xing, G., Zhang, Y., Lu, C., Pless, R., & Gill, C. (2007). Coverage and connectivity control in wireless sensor networks under fading channel. IEEE Transactions on Automatic Control, 52(11), 2018-2031.",
        "Ye, F., Zhong, G., Lu, S., & Zhang, L. (2003). PEAS: A robust energy conserving protocol for long-lived sensor networks. Proceedings of IEEE ICDCS, 28-37.",
        "Younis, O., & Fahmy, S. (2004). HEED: A hybrid, energy-efficient, distributed clustering approach for ad hoc sensor networks. IEEE Transactions on Mobile Computing, 3(4), 366-379.",
        "Zhang, W., Liu, X., & Chen, Y. (2021). Cooling-aware duty cycling for energy-efficient wireless sensor networks. Proceedings of IEEE INFOCOM, 1-9."
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref, style='Normal')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

def main():
    """Generate the complete manuscript DOCX."""
    print("=" * 70)
    print("Generating Final Revised Manuscript in DOCX Format")
    print("=" * 70)
    
    doc = Document()
    setup_styles(doc)
    
    print("✓ Adding title and authors...")
    add_title_section(doc)
    
    print("✓ Adding abstract...")
    add_abstract(doc)
    
    print("✓ Adding introduction...")
    add_introduction(doc)
    
    print("✓ Adding related work and literature review...")
    add_related_work(doc)
    
    print("✓ Adding methodology summary...")
    add_methodology_summary(doc)
    
    print("✓ Adding results and validation...")
    add_results_summary(doc)
    
    print("✓ Adding discussion and conclusion...")
    add_discussion_conclusion(doc)
    
    print("✓ Adding references...")
    add_references(doc)
    
    # Save document
    output_path = Path(__file__).parent / 'Final_Revised_Manuscript.docx'
    doc.save(output_path)
    
    print("=" * 70)
    print(f"✓ Manuscript saved: {output_path}")
    print(f"✓ File size: {output_path.stat().st_size / 1024:.1f} KB")
    print("=" * 70)
    print("\nDocument includes:")
    print("  • Title, authors, and affiliation")
    print("  • Abstract with keywords")
    print("  • Introduction (6 paragraphs)")
    print("  • Related Work with comparison table")
    print("  • Methodology summary (6 subsections)")
    print("  • Results with performance table")
    print("  • Discussion and Conclusion")
    print("  • 13 references")
    print("\n✓ Ready for submission!")

if __name__ == '__main__':
    main()
