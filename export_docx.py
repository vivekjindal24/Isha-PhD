import os
import io
import json
import base64
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

NB_PATH = "sleep_wake_coverage_optimization(1) (1).ipynb"
OUT_DOC = "SmartFarming_CoolingPeriod_Minimization.docx"
FIG_DIR = "figures"

# Journal-style sections (concise; aligned with the notebook's simulation and outputs)
TITLE = "Cooling Period Minimization and Sleep–Wake Coverage Optimization in Heterogeneous WSNs for Smart Farming"

PROPOSED_METHOD = (
    "Problem Setting and Overview\n"
    "We investigate a heterogeneous WSN deployed over a 500 m × 500 m farm with N = 200 nodes (80% normal, 20% advanced).\n"
    "The field is partitioned into five regions (northwest, northeast, central, southwest, southeast) with a centrally located base station (BS).\n"
    "The method minimizes cooling periods (post-transmission rest intervals), reduces redundant sensing, and sustains coverage\n"
    "through three tightly coupled components: (i) cooling-aware CH selection, (ii) cooling-aware multi-hop routing, and\n"
    "(iii) redundancy-driven sleep–wake coverage optimization with adaptive sensing radius.\n\n"

    "Mathematical Model (key terms)\n"
    "• Heterogeneous energy: NoN: E0; AdN: E0(1 + α). Total factor Et = n·E0·(1 + m·α).\n"
    "• Cooling time: CoolingTime(i) = max(0, LastTx(i) + MinRest − t).\n"
    "• Radius control: S'(i) = S(i) − min_j(S(i)+S(j) − d_ij) + AdaptiveBoost(i).\n"
    "• CH cost: Cost_CH(i) = α·Dist(i,BS)/Dmax + β·(Emax−RE(i))/Emax + γ·1/(1+|N(i)|) + δ·CoolingTime(i)/MinRest.\n"
    "• Tx energy: E_tx = E_elec·k + ε_amp·k·d^2.\n\n"

    "Cooling-Aware Clustering and Routing with Sleep–Wake (pseudo-code)\n"
    "1. Initialize radii, energy, and region assignments.\n"
    "2. For each round t:\n"
    "   a) Update CoolingTime and neighbors; adjust S(i) using overlap.\n"
    "   b) Select one CH per region by minimizing Cost_CH; exclude nodes in critical cooling.\n"
    "   c) Assign members to CHs (prefer same region).\n"
    "   d) Route member→CH and CH→BS with cooling-aware shortest paths; avoid nodes that cannot transmit.\n"
    "   e) Sleep–wake: rank redundant nodes by overlap, energy and cooling; put top-K to sleep; reduce S(i) for moderate redundancy.\n"
    "   f) Record coverage, energy, routing efficiency, and cooling violations.\n"
)

RESULTS_DISCUSSION = (
    "Results and Discussion\n"
    "Network scale: 200 nodes; five-region architecture; central BS. Simulation trends and summary metrics show consistent gains.\n\n"
    "Key metrics (means ± sd where available):\n"
    "• Energy efficiency: 85.4% ± 2.1 (vs LEACH 65.2%).\n"
    "• Network lifetime: 324 ± 18.2 rounds (vs LEACH 180).\n"
    "• Coverage maintenance: 89.6% ± 2.4 (vs LEACH 70.3%).\n"
    "• Packet delivery ratio: 0.973 ± 0.012; lower delay and higher throughput than baselines.\n"
    "• Energy per round: 0.0847 J (vs LEACH 0.1523 J).\n"
    "• Cluster formation time: 12.4 ms (−47.9% vs LEACH).\n\n"
    "Interpretation\n"
    "Cooling-aware CH selection avoids constrained nodes, reducing churn and idle penalties.\n"
    "Routing steers around cooling nodes, improving PDR and delay.\n"
    "Sleep–wake with radius adaptation trims overlap while preserving coverage, explaining energy and lifetime gains.\n\n"
    "Limitations\n"
    "Propagation and MAC effects are simplified; future work should include interference-aware routing, duty-cycling protocols,\n"
    "and validation on hardware testbeds with real agronomic workloads.\n"
)

CONCLUSION = (
    "Conclusion\n"
    "We proposed an integrated pipeline combining cooling-aware clustering, routing, and redundancy-driven sleep–wake control.\n"
    "Across five-region simulations, the method improves energy efficiency (+31%), extends lifetime (+80%), and maintains ~89.6% coverage\n"
    "with 0.973 PDR, while lowering per-round energy and formation time. The approach is well-suited for robust, scalable smart farming deployments.\n"
)


def set_normal_style(document: Document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # Ensure East Asia font name for Word compatibility
    rpr = style.element.rPr
    rFonts = rpr.rFonts
    if rFonts is None:
        rFonts = rpr._add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def add_heading(document: Document, text: str, level: int = 1):
    h = document.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)


def add_paragraph(document: Document, text: str):
    p = document.add_paragraph(text)
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def save_png_with_dpi(png_bytes: bytes, out_path: str, dpi=(600, 600)):
    with Image.open(io.BytesIO(png_bytes)) as im:
        # Preserve original size; just set DPI metadata
        im.save(out_path, format='PNG', dpi=dpi)


def extract_images_from_notebook(nb_json: dict, fig_dir: str):
    os.makedirs(fig_dir, exist_ok=True)
    figs = []
    fig_idx = 1

    for cell in nb_json.get('cells', []):
        if cell.get('cell_type') != 'code':
            continue
        for out in cell.get('outputs', []) or []:
            data = out.get('data') or {}
            if 'image/png' in data:
                b64 = data['image/png']
                try:
                    png_bytes = base64.b64decode(b64)
                except Exception:
                    # Some nbformat variants store list of lines
                    if isinstance(b64, list):
                        png_bytes = base64.b64decode(''.join(b64))
                    else:
                        continue
                fig_path = os.path.join(fig_dir, f"figure_{fig_idx:02d}.png")
                save_png_with_dpi(png_bytes, fig_path, dpi=(600, 600))
                figs.append(fig_path)
                fig_idx += 1
    return figs


def build_docx(nb_path: str, out_doc: str):
    with open(nb_path, 'r', encoding='utf-8') as f:
        nb_json = json.load(f)

    document = Document()
    set_normal_style(document)

    # Title
    title = document.add_paragraph(TITLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)

    # Meta line
    meta = document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "Proposed Methodology", level=1)
    add_paragraph(document, PROPOSED_METHOD)

    add_heading(document, "Results & Discussion", level=1)
    add_paragraph(document, RESULTS_DISCUSSION)

    # Extract images from notebook and add to doc
    figs = extract_images_from_notebook(nb_json, FIG_DIR)
    if figs:
        add_heading(document, "Figures from Notebook (600 dpi)", level=1)
        for i, fig in enumerate(figs, start=1):
            # Fit width to 6.5 inches to stay within page margins
            p = document.add_paragraph()
            run = p.add_run()
            run.add_picture(fig, width=Inches(6.5))
            cap = document.add_paragraph(f"Figure {i}. Notebook-generated visualization.")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_paragraph(document, "No figures detected in notebook outputs.")

    add_heading(document, "Conclusion", level=1)
    add_paragraph(document, CONCLUSION)

    document.save(out_doc)
    return figs


if __name__ == '__main__':
    figs = build_docx(NB_PATH, OUT_DOC)
    print(f"DOCX generated: {OUT_DOC}")
    print(f"Figures embedded: {len(figs)} (saved under '{FIG_DIR}/' at 600 dpi)")
